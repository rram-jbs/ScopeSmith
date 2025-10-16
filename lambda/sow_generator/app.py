import os
import json
import boto3
from datetime import datetime
import tempfile
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def handler(event, context):
    try:
        print(f"[SOW GENERATOR] Received event: {json.dumps(event)}")
        
        # Check if this is a Bedrock Agent action group invocation
        if 'agent' in event or 'actionGroup' in event or 'function' in event:
            print(f"[SOW GENERATOR] Bedrock Agent action group invocation detected")
            
            parameters = event.get('parameters', [])
            param_dict = {}
            for param in parameters:
                param_name = param['name']
                param_value = param['value']
                
                # Handle nested JSON in parameters
                if param_name in ['proposal_data']:
                    try:
                        if isinstance(param_value, str):
                            param_dict[param_name] = json.loads(param_value)
                        else:
                            param_dict[param_name] = param_value
                    except json.JSONDecodeError as e:
                        print(f"[SOW GENERATOR] JSON decode error for {param_name}: {str(e)}")
                        param_dict[param_name] = param_value
                else:
                    param_dict[param_name] = param_value
            
            session_id = param_dict.get('session_id')
            template_path = param_dict.get('template_path')
            proposal_data = param_dict.get('proposal_data')
        elif 'inputText' in event:
            try:
                params = json.loads(event['inputText'])
                session_id = params['session_id']
                template_path = params.get('template_path')
                proposal_data = params.get('proposal_data')
                if isinstance(proposal_data, str):
                    proposal_data = json.loads(proposal_data)
            except (json.JSONDecodeError, KeyError) as e:
                return {
                    'statusCode': 400,
                    'body': json.dumps({
                        'error': f'Invalid input format from AgentCore: {str(e)}'
                    })
                }
        else:
            return format_agent_response(400, 'Missing required parameters')

        print(f"[SOW GENERATOR] Generating SOW document for session: {session_id}")
        
        # Initialize AWS clients
        s3 = boto3.client('s3')
        dynamodb = boto3.client('dynamodb')
        bedrock = boto3.client('bedrock-runtime')
        
        # Update status
        dynamodb.update_item(
            TableName=os.environ['SESSIONS_TABLE_NAME'],
            Key={'session_id': {'S': session_id}},
            UpdateExpression='SET #status = :status, updated_at = :ua',
            ExpressionAttributeNames={'#status': 'status'},
            ExpressionAttributeValues={
                ':status': {'S': 'GENERATING_SOW'},
                ':ua': {'S': datetime.utcnow().isoformat()}
            }
        )
        
        # Get complete session data including raw requirements
        response = dynamodb.get_item(
            TableName=os.environ['SESSIONS_TABLE_NAME'],
            Key={'session_id': {'S': session_id}}
        )
        
        if 'Item' not in response:
            raise ValueError(f"Session {session_id} not found")
        
        session_data = response['Item']
        requirements_data_str = session_data.get('requirements_data', {}).get('S', '{}')
        cost_data_str = session_data.get('cost_data', {}).get('S', '{}')
        
        try:
            requirements_data_full = json.loads(requirements_data_str)
            raw_requirements = requirements_data_full.get('raw_requirements', '')
            analysis_data = requirements_data_full.get('analysis', {})
            cost_data = json.loads(cost_data_str)
        except json.JSONDecodeError:
            raw_requirements = ''
            analysis_data = {}
            cost_data = {}
        
        # Use Bedrock to generate SOW content with full context
        prompt = f"""Generate a comprehensive Statement of Work (SOW) document based on the following project information.

RAW PROJECT REQUIREMENTS (Original User Input):
{raw_requirements}

ANALYZED PROJECT DATA:
{json.dumps(analysis_data, indent=2)}

COST DATA:
{json.dumps(cost_data, indent=2)}

Generate a professional SOW with the following sections. Return as JSON (no markdown, no code blocks):
{{
    "project_overview": "Detailed project overview",
    "scope_of_services": ["List of services to be provided"],
    "deliverables": ["List of specific deliverables with descriptions"],
    "timeline": "Project timeline with phases",
    "costs": {{"item_name": {{"description": "desc", "amount": 0}}}},
    "assumptions": ["Key assumptions"],
    "exclusions": ["What is not included"],
    "acceptance_criteria": ["Criteria for acceptance"]
}}"""
        
        print(f"[SOW GENERATOR] Calling Bedrock for SOW content generation...")
        
        bedrock_response = bedrock.invoke_model(
            modelId=os.environ['BEDROCK_MODEL_ID'],
            contentType='application/json',
            body=json.dumps({
                "messages": [
                    {
                        "role": "user",
                        "content": [{"text": prompt}]
                    }
                ],
                "inferenceConfig": {
                    "max_new_tokens": 3000,
                    "temperature": 0.7
                }
            })
        )
        
        result = json.loads(bedrock_response['body'].read().decode())
        
        # Parse the response - handle different response structures
        sow_content = None
        if 'content' in result and len(result['content']) > 0:
            sow_content = result['content'][0].get('text', '')
        elif 'output' in result:
            if isinstance(result['output'], dict) and 'message' in result['output']:
                message = result['output']['message']
                if 'content' in message and len(message['content']) > 0:
                    sow_content = message['content'][0].get('text', '')
        
        if not sow_content:
            print(f"[SOW GENERATOR] WARNING: Could not extract text from model response")
            # Use fallback data
            proposal_data = {
                'project_overview': analysis_data.get('project_scope', 'Project overview'),
                'services': analysis_data.get('deliverables', []),
                'timeline': analysis_data.get('timeline_estimate', 'TBD'),
                'deliverables': analysis_data.get('deliverables', []),
                'costs': {}
            }
        else:
            # Parse JSON from response
            try:
                if '```json' in sow_content:
                    sow_content = sow_content.split('```json')[1].split('```')[0].strip()
                elif '```' in sow_content:
                    sow_content = sow_content.split('```')[1].split('```')[0].strip()
                
                proposal_data = json.loads(sow_content)
            except json.JSONDecodeError as e:
                print(f"[SOW GENERATOR] Could not parse SOW content as JSON: {str(e)}")
                # Use fallback
                proposal_data = {
                    'project_overview': analysis_data.get('project_scope', 'Project overview'),
                    'services': analysis_data.get('deliverables', []),
                    'timeline': analysis_data.get('timeline_estimate', 'TBD'),
                    'deliverables': analysis_data.get('deliverables', []),
                    'costs': {}
                }
        
        # Download template or create new document
        doc = None
        if template_path and not template_path.startswith('path/to/'):
            # Check if template exists in S3 before downloading
            try:
                s3.head_object(
                    Bucket=os.environ['TEMPLATES_BUCKET_NAME'],
                    Key=template_path
                )
                # Template exists, download it
                with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as template_file:
                    s3.download_file(
                        os.environ['TEMPLATES_BUCKET_NAME'],
                        template_path,
                        template_file.name
                    )
                    doc = Document(template_file.name)
                    print(f"[SOW GENERATOR] Using template: {template_path}")
            except s3.exceptions.NoSuchKey:
                print(f"[SOW GENERATOR] Template not found: {template_path}, creating blank document")
                doc = Document()
            except Exception as e:
                print(f"[SOW GENERATOR] Error loading template: {str(e)}, creating blank document")
                doc = Document()
        else:
            print(f"[SOW GENERATOR] No valid template specified, creating blank document")
            doc = Document()
        
        # Add SOW content to document
        doc.add_heading('Statement of Work', 0)
        
        doc.add_heading('1. Project Overview', level=1)
        doc.add_paragraph(proposal_data.get('project_overview', ''))
        
        doc.add_heading('2. Scope of Services', level=1)
        services = proposal_data.get('scope_of_services', proposal_data.get('services', []))
        for service in services:
            doc.add_paragraph(service, style='List Bullet')
        
        doc.add_heading('3. Timeline', level=1)
        doc.add_paragraph(proposal_data.get('timeline', ''))
        
        doc.add_heading('4. Deliverables', level=1)
        deliverables = proposal_data.get('deliverables', [])
        for deliverable in deliverables:
            doc.add_paragraph(str(deliverable), style='List Number')
        
        doc.add_heading('5. Cost Breakdown', level=1)
        costs = proposal_data.get('costs', {})
        if costs:
            table = doc.add_table(rows=1, cols=3)
            table.style = 'Table Grid'
            header_cells = table.rows[0].cells
            header_cells[0].text = 'Item'
            header_cells[1].text = 'Description'
            header_cells[2].text = 'Cost'
            
            for item, details in costs.items():
                row_cells = table.add_row().cells
                row_cells[0].text = item
                row_cells[1].text = details.get('description', '')
                amount = details.get('amount', 0)
                row_cells[2].text = f"${amount:,.2f}" if isinstance(amount, (int, float)) else str(amount)
        
        # Add assumptions if present
        if 'assumptions' in proposal_data:
            doc.add_heading('6. Assumptions', level=1)
            for assumption in proposal_data.get('assumptions', []):
                doc.add_paragraph(assumption, style='List Bullet')
        
        # Add exclusions if present
        if 'exclusions' in proposal_data:
            doc.add_heading('7. Exclusions', level=1)
            for exclusion in proposal_data.get('exclusions', []):
                doc.add_paragraph(exclusion, style='List Bullet')
        
        # Save modified document
        output_path = f"{session_id}/sow.docx"
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as output_file:
            doc.save(output_file.name)
            
            # Upload to S3
            s3.upload_file(
                output_file.name,
                os.environ['ARTIFACTS_BUCKET_NAME'],
                output_path
            )
        
        # Generate presigned URL
        presigned_url = s3.generate_presigned_url(
            'get_object',
            Params={
                'Bucket': os.environ['ARTIFACTS_BUCKET_NAME'],
                'Key': output_path
            },
            ExpiresIn=3600
        )
        
        # Update session with document URL
        dynamodb.update_item(
            TableName=os.environ['SESSIONS_TABLE_NAME'],
            Key={'session_id': {'S': session_id}},
            UpdateExpression='SET document_urls = list_append(if_not_exists(document_urls, :empty), :url), #status = :status, updated_at = :ua',
            ExpressionAttributeNames={'#status': 'status'},
            ExpressionAttributeValues={
                ':url': {'L': [{'S': presigned_url}]},
                ':empty': {'L': []},
                ':status': {'S': 'SOW_GENERATED'},
                ':ua': {'S': datetime.utcnow().isoformat()}
            }
        )
        
        print(f"[SOW GENERATOR] SOW generated successfully")
        
        return format_agent_response(200, {
            'session_id': session_id,
            'message': 'SOW generated successfully',
            'document_url': presigned_url
        })
        
    except Exception as e:
        print(f"[SOW GENERATOR] ERROR: {str(e)}")
        import traceback
        print(f"[SOW GENERATOR] Traceback: {traceback.format_exc()}")
        return format_agent_response(500, f'SOW generation failed: {str(e)}')

def format_agent_response(status_code, body):
    """Format response for Bedrock Agent action group"""
    if isinstance(body, dict):
        body_text = json.dumps(body)
    else:
        body_text = str(body)
    
    return {
        'messageVersion': '1.0',
        'response': {
            'actionGroup': 'SOWGenerator',
            'function': 'sowgenerator',  # Must match function name in action group
            'functionResponse': {
                'responseBody': {
                    'TEXT': {
                        'body': body_text
                    }
                }
            }
        }
    }